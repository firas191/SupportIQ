"""Extraction de texte depuis des documents non structurés (S7-J4, rapport §5.4).

```
PDF natif ──PyMuPDF──> texte
PDF scanné ──PyMuPDF (rendu) ──> images ──Tesseract──> texte
DOCX ──python-docx──> texte
TXT / MD ──décodage tolérant──> texte
```

**L'OCR est un repli, jamais le chemin principal.** C'est le plan B du rapport §11, et il est
appliqué à la lettre : on tente d'abord l'extraction native, qui est exacte, rapide et gratuite ; on
ne bascule sur Tesseract que si le document ne rend visiblement pas de texte. Faire l'inverse
coûterait des secondes par page et introduirait des fautes de reconnaissance dans des documents qui
n'en avaient aucune.

**Comment on détecte un PDF scanné.** Pas par un drapeau — il n'y en a pas. Par le rapport
*caractères extraits / nombre de pages* : un PDF natif rend typiquement plus d'un millier de
caractères par page, un PDF image en rend zéro ou quelques dizaines (les artefacts de métadonnées).
Le seuil est franc, l'ambiguïté rare, et un faux positif ne coûte qu'un passage OCR inutile.

Tous les imports lourds sont **paresseux et résilients** : sans python-docx seuls les DOCX
échouent, sans Tesseract seuls les PDF scannés échouent — avec un message qui dit lequel manque, et
non une erreur 500 opaque.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass

from app.kb.loader import UnsupportedDocument, decode_text, extension_of, normalise

logger = logging.getLogger(__name__)

PDF_EXT = {".pdf"}
DOCX_EXT = {".docx"}
TEXT_EXT = {".txt", ".md", ".markdown"}
SUPPORTED_EXT = PDF_EXT | DOCX_EXT | TEXT_EXT

#: En dessous de ce nombre de caractères par page, le PDF est considéré comme une image.
#: Une page de texte dense en contient 2 000 à 3 000 ; une page scannée en rend 0 à quelques
#: dizaines. Le seuil est volontairement bas : basculer en OCR à tort ne coûte que du temps,
#: alors que ne pas basculer produit un document vide et un lot de tickets fantômes.
SCANNED_CHARS_PER_PAGE = 80

#: Résolution de rendu avant OCR. 200 dpi est le plancher usuel pour Tesseract sur du corps de
#: texte ; en dessous les caractères se collent, au-dessus le gain est marginal et la mémoire
#: grimpe au carré.
OCR_DPI = 200


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    pages: int
    #: `native`, `ocr` ou `plain`. Voyage jusqu'à l'interface : un texte issu d'OCR mérite une
    #: relecture plus attentive, et l'agent doit savoir lequel il relit.
    method: str


def extract(filename: str, data: bytes) -> ExtractedDocument:
    ext = extension_of(filename)
    if ext not in SUPPORTED_EXT:
        raise UnsupportedDocument(
            f"Format {ext or 'inconnu'} non pris en charge (attendus : "
            f"{', '.join(sorted(SUPPORTED_EXT))})"
        )

    if ext in PDF_EXT:
        text, pages, method = _read_pdf(data)
    elif ext in DOCX_EXT:
        text, pages, method = _read_docx(data), 1, "native"
    else:
        text, pages, method = decode_text(data), 1, "plain"

    text = normalise(text)
    if not text:
        raise UnsupportedDocument("Le document ne contient aucun texte exploitable")

    return ExtractedDocument(text=text, pages=pages, method=method)


def _read_docx(data: bytes) -> str:
    """Paragraphes **et** cellules de tableau.

    Ne lire que `document.paragraphs` est l'erreur classique : une demande client mise en forme
    dans un tableau — ce que produit tout formulaire exporté en Word — ressortirait entièrement
    vide, sans aucune erreur pour le signaler.
    """
    try:
        import io

        import docx
    except ImportError as exc:  # pragma: no cover - depend de l'environnement
        raise UnsupportedDocument(
            "Lecture DOCX indisponible sur ce serveur (python-docx absent)"
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        # Pas de `noqa: BLE001` ici : la regle vise les `except` qui **avalent** l'exception, or
        # celle-ci est immediatement retraduite et relancee.
        raise UnsupportedDocument(f"DOCX illisible : {exc}") from exc

    blocks = [p.text.strip() for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(b for b in blocks if b)


def _read_pdf(data: bytes) -> tuple[str, int, str]:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - depend de l'environnement
        raise UnsupportedDocument("Lecture PDF indisponible sur ce serveur (PyMuPDF absent)") from exc

    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            pages = len(doc)
            # `sort=True` : ordre de lecture visuel, pas ordre du flux interne (leçon du S5-J1 —
            # sans lui, un document sur deux colonnes sort entrelacé).
            native = "\n\n".join(
                page.get_text("text", sort=True).strip() for page in doc
            ).strip()

            if pages and len(native) / pages >= SCANNED_CHARS_PER_PAGE:
                return native, pages, "native"

            logger.info(
                "PDF probablement scanne (%d caracteres pour %d pages) - passage en OCR",
                len(native), pages,
            )
            ocr = _ocr_pdf(doc)
            # Si l'OCR rend moins que l'extraction native, on garde la native : mieux vaut peu de
            # texte exact qu'un peu plus de texte inventé par une reconnaissance qui a échoué.
            return (ocr, pages, "ocr") if len(ocr) > len(native) else (native, pages, "native")

    except UnsupportedDocument:
        raise
    except Exception as exc:
        raise UnsupportedDocument(f"PDF illisible : {exc}") from exc


def _ocr_pdf(doc) -> str:
    """OCR page par page. Renvoie une chaîne vide si Tesseract est absent."""
    try:
        import io

        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("OCR indisponible (pytesseract/Pillow absents) - PDF scanne non exploitable")
        return ""

    texts: list[str] = []
    for index, page in enumerate(doc):
        try:
            pixmap = page.get_pixmap(dpi=OCR_DPI)
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            # `fra+eng` : le corpus est bilingue, et Tesseract choisit mal quand on ne lui donne
            # qu'une langue. Si le pack français manque, l'appel échoue — d'où le try par page,
            # qui laisse les autres pages aboutir.
            texts.append(pytesseract.image_to_string(image, lang="fra+eng"))
        except Exception as exc:  # noqa: BLE001
            logger.warning("OCR echoue sur la page %d : %s", index + 1, exc)

    return "\n\n".join(t.strip() for t in texts if t.strip())
